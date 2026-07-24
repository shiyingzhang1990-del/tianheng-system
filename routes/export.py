"""
导出功能路由
"""
from flask import Blueprint, request, jsonify, send_file, make_response
from datetime import datetime
from urllib.parse import quote
import io
import json

from models import db
from models.qa_record import QARecord

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx未安装，Word导出功能将不可用")

# 创建蓝图
export_bp = Blueprint('export', __name__)

# 默认用户ID（无认证模式）
DEFAULT_USER_ID = 1


@export_bp.route('/qa/<int:qa_id>/markdown', methods=['GET'])
def export_qa_markdown(qa_id):
    """导出单条问答为Markdown格式"""
    try:
        # 获取问答记录
        qa_record = QARecord.query.filter_by(
            id=qa_id,
            user_id=DEFAULT_USER_ID
        ).first()
        
        if not qa_record:
            return jsonify({'error': '问答记录不存在'}), 404
        
        # 构建Markdown内容
        markdown_content = f"""# AI问答记录

## 问题
{qa_record.question}

## 推理过程
{qa_record.reasoning if qa_record.reasoning else '（无推理过程）'}

## 回答
{qa_record.answer}

---

**响应时间**: {qa_record.response_time}秒  
**生成时间**: {qa_record.created_time.strftime('%Y-%m-%d %H:%M:%S')}  
**来源**: 天衡系统（壬镜科技）
"""
        
        # 创建文件名（使用URL编码支持中文）
        filename = f"问答_{qa_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        encoded_filename = quote(filename)
        
        # 创建响应
        response = make_response(markdown_content)
        response.headers['Content-Type'] = 'text/markdown; charset=utf-8'
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        
        return response
        
    except Exception as e:
        print(f"导出Markdown失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'导出失败: {str(e)}'}), 500


@export_bp.route('/qa/<int:qa_id>/word', methods=['GET'])
def export_qa_word(qa_id):
    """导出单条问答为Word文档（不包含推理过程）"""
    try:
        if not DOCX_AVAILABLE:
            return jsonify({'error': 'Word导出功能未启用，请安装python-docx'}), 500
        
        # 获取问答记录
        qa_record = QARecord.query.filter_by(
            id=qa_id,
            user_id=DEFAULT_USER_ID
        ).first()
        
        if not qa_record:
            return jsonify({'error': '问答记录不存在'}), 404
        
        # 创建Word文档
        doc = DocxDocument()
        
        # 设置文档标题
        title = doc.add_heading('AI问答记录', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加元信息
        info_para = doc.add_paragraph()
        info_para.add_run(f'生成时间：{qa_record.created_time.strftime("%Y-%m-%d %H:%M:%S")}').font.size = Pt(10)
        info_para.add_run(f'\n响应时间：{qa_record.response_time}秒').font.size = Pt(10)
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()  # 空行
        
        # 添加问题
        question_heading = doc.add_heading('问题', 1)
        question_heading.runs[0].font.color.rgb = RGBColor(124, 58, 237)  # 紫色
        
        question_para = doc.add_paragraph(qa_record.question)
        question_para.paragraph_format.space_after = Pt(12)
        
        # 添加回答（不包含推理过程）
        answer_heading = doc.add_heading('回答', 1)
        answer_heading.runs[0].font.color.rgb = RGBColor(124, 58, 237)  # 紫色
        
        # 处理回答内容（支持多段）
        answer_lines = qa_record.answer.split('\n')
        for line in answer_lines:
            if line.strip():
                doc.add_paragraph(line.strip())
        
        # 添加页脚信息
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run('——————————————————————').font.color.rgb = RGBColor(200, 200, 200)
        footer_info = doc.add_paragraph()
        footer_info.add_run('来源：天衡系统（壬镜科技）').font.size = Pt(9)
        footer_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # 创建文件名
        filename = f"问答_{qa_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        encoded_filename = quote(filename)
        
        # 返回文件
        return send_file(
            file_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        
    except Exception as e:
        print(f"导出Word失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'导出失败: {str(e)}'}), 500


@export_bp.route('/qa/<int:qa_id>/text', methods=['GET'])
def export_qa_text(qa_id):
    """导出单条问答为纯文本格式"""
    try:
        qa_record = QARecord.query.filter_by(
            id=qa_id,
            user_id=DEFAULT_USER_ID
        ).first()
        
        if not qa_record:
            return jsonify({'error': '问答记录不存在'}), 404
        
        # 构建纯文本内容
        text_content = f"""AI问答记录
{'=' * 60}

问题：
{qa_record.question}

推理过程：
{qa_record.reasoning if qa_record.reasoning else '（无推理过程）'}

回答：
{qa_record.answer}

{'=' * 60}
响应时间: {qa_record.response_time}秒
生成时间: {qa_record.created_time.strftime('%Y-%m-%d %H:%M:%S')}
来源: 天衡系统（壬镜科技）
"""
        
        filename = f"问答_{qa_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        encoded_filename = quote(filename)
        
        response = make_response(text_content)
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        
        return response
        
    except Exception as e:
        print(f"导出文本失败: {e}")
        return jsonify({'error': f'导出失败: {str(e)}'}), 500


@export_bp.route('/qa/<int:qa_id>/json', methods=['GET'])
def export_qa_json(qa_id):
    """导出单条问答为JSON格式"""
    try:
        qa_record = QARecord.query.filter_by(
            id=qa_id,
            user_id=DEFAULT_USER_ID
        ).first()
        
        if not qa_record:
            return jsonify({'error': '问答记录不存在'}), 404
        
        # 构建JSON数据
        data = {
            'id': qa_record.id,
            'question': qa_record.question,
            'reasoning': qa_record.reasoning,
            'answer': qa_record.answer,
            'sources': qa_record.sources,
            'response_time': qa_record.response_time,
            'satisfaction': qa_record.satisfaction,
            'created_time': qa_record.created_time.strftime('%Y-%m-%d %H:%M:%S'),
            'export_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'source_system': '天衡系统（壬镜科技）'
        }
        
        filename = f"问答_{qa_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        encoded_filename = quote(filename)
        
        response = make_response(json.dumps(data, ensure_ascii=False, indent=2))
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        
        return response
        
    except Exception as e:
        print(f"导出JSON失败: {e}")
        return jsonify({'error': f'导出失败: {str(e)}'}), 500


@export_bp.route('/history/markdown', methods=['POST'])
def export_history_markdown():
    """导出多条问答历史为Markdown格式"""
    try:
        data = request.get_json()
        qa_ids = data.get('qa_ids', [])
        
        if not qa_ids:
            return jsonify({'error': '请选择要导出的问答记录'}), 400
        
        # 获取问答记录
        qa_records = QARecord.query.filter(
            QARecord.id.in_(qa_ids),
            QARecord.user_id == DEFAULT_USER_ID
        ).order_by(QARecord.created_time.desc()).all()
        
        if not qa_records:
            return jsonify({'error': '未找到问答记录'}), 404
        
        # 构建Markdown内容
        markdown_content = f"""# AI问答历史记录

**导出时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**记录数量**: {len(qa_records)}条

---

"""
        
        for idx, record in enumerate(qa_records, 1):
            markdown_content += f"""
## {idx}. {record.question}

**提问时间**: {record.created_time.strftime('%Y-%m-%d %H:%M:%S')}  
**响应时间**: {record.response_time}秒

### 推理过程
{record.reasoning if record.reasoning else '（无推理过程）'}

### 回答
{record.answer}

---

"""
        
        filename = f"问答历史_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        encoded_filename = quote(filename)
        
        response = make_response(markdown_content)
        response.headers['Content-Type'] = 'text/markdown; charset=utf-8'
        response.headers['Content-Disposition'] = f"attachment; filename*=UTF-8''{encoded_filename}"
        
        return response
        
    except Exception as e:
        print(f"导出历史记录失败: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'导出失败: {str(e)}'}), 500